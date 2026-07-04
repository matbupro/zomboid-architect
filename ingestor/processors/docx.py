"""
docx — Processeur pour fichiers Word (.docx).

Extraction du texte brut des paragraphs, tables, et sections.
"""

from __future__ import annotations

import time
from pathlib import Path

from src.governance.logger import get_logger

logger = get_logger(__name__)


class DocxProcessor:
    """Extracteur de contenu Word .docx."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self, config):
        self.config = config

    async def extract(self, source: str) -> "ExtractionResult":  # type: ignore[name-defined] # noqa: F821
        """Extrait le texte d'un fichier .docx.

        Returns:
            ExtractionResult avec chunks du contenu Word.
        """
        from pathlib import Path
        from .base import Chunk, ExtractionResult  # local import pour éviter circular dep

        p = Path(source)
        start_time = time.monotonic()

        if not p.exists():
            raise FileNotFoundError(f"Fichier Word non trouvé : {p}")

        try:
            import docx as python_docx  # type: ignore[import-not-found]
            doc = python_docx.Document(str(p))

            # Extraction du texte des paragraphs + tables
            text_parts = []
            table_num = 0

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                table_text = f"[TABLE {table_num}]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_text += row_text + "\n"
                text_parts.append(table_text)
                table_num += 1

            full_text = "\n".join(text_parts)

        except ImportError:
            logger.error("python-docx non installé : pip install python-docx")
            raise

        if not full_text.strip():
            return ExtractionResult(
                chunks=[], source=str(source), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_hash="", word_count=0, extraction_time_ms=(time.monotonic() - start_time) * 1000,
                metadata={"error": "Aucun texte trouvé dans le document Word"},
            )

        chunks = [Chunk(text=full_text[:self.config.CHUNK_SIZE], index=0, start_offset=0)]
        if len(full_text) > self.config.CHUNK_SIZE:
            remaining = full_text[self.config.CHUNK_SIZE:]
            chunks.append(Chunk(text=remaining[:self.config.CHUNK_SIZE], index=1, start_offset=self.config.CHUNK_SIZE))

        word_count = len(full_text.split())
        duration_ms = (time.monotonic() - start_time) * 1000

        return ExtractionResult(
            chunks=chunks,
            source=str(source),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_hash=self._compute_file_hash(p),
            word_count=word_count,
            extraction_time_ms=duration_ms,
            metadata={
                "paragraphs": len(doc.paragraphs) if 'doc' in locals() else 0,
                "tables": table_num,
                "file_size_bytes": p.stat().st_size,
            },
        )

    def _compute_file_hash(self, path: Path) -> str:
        """SHA-256 du contenu brut."""
        import hashlib
        h = hashlib.sha256()
        try:
            with open(path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    h.update(chunk)
        except (OSError, IOError):
            pass
        return h.hexdigest()
